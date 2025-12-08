import streamlit as st
import pandas as pd
from datetime import datetime
import os
from io import BytesIO
import base64
import tempfile
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit.components.v1 import html as st_html
from utils.quotation_utils import render_quotation_html, html_to_pdf
from pathlib import Path
import sys
import threading
import json
import subprocess
import time
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from utils.logger import log_event
from utils.settings import load_settings
try:
    from utils import db as _db
except Exception:
    _db = None
 # تم حذف الاستيراد غير المستخدم requests

def proper_case(text):
    if not text:
        return ""
    try:
        return text.title().strip()
    except:
        return text

# Apply the same visual theme used in dashboard_page.py
def _apply_quotation_theme():
    # Now inherits global Invoice theme from main.py
    st.markdown("<style></style>", unsafe_allow_html=True)

def quotation_app():
    _apply_quotation_theme()

        # (Header hero removed to match invoice page)

    # =========================
    # Helper
    # =========================
    def format_phone_input(raw_input):
        digits = ''.join(filter(str.isdigit, raw_input))
        if digits.startswith("0"): digits = digits[1:]
        if digits.startswith("5") and len(digits) == 9:
            return f"+971 {digits[:2]} {digits[2:5]} {digits[5:]}"
        return None

    # =========================
    # UAE Locations
    # =========================
    uae_locations = [
        "Abu Dhabi - Al Shamkha","Abu Dhabi - Al Shawamekh","Abu Dhabi - Khalifa City",
        "Abu Dhabi - Al Bateen","Abu Dhabi - Al Reem Island","Abu Dhabi - Yas Island",
        "Abu Dhabi - Al Mushrif","Abu Dhabi - Al Rawdah","Abu Dhabi - Al Muroor",
        "Abu Dhabi - Baniyas","Abu Dhabi - Mussafah","Abu Dhabi - Al Mafraq",
        "Abu Dhabi - Al Falah","Abu Dhabi - MBZ City","Abu Dhabi - Al Raha",
        "Abu Dhabi - Al Maqtaa","Abu Dhabi - Zayed Port","Abu Dhabi - Saadiyat Island",
        "Al Ain - Al Jimi","Al Ain - Falaj Hazza","Al Ain - Al Maqam",
        "Al Ain - Zakher","Al Ain - Hili","Al Ain - Al Foah","Al Ain - Al Mutaredh",
        "Al Ain - Al Towayya","Al Ain - Al Sarooj","Al Ain - Al Nyadat",
        "Dubai - Marina","Dubai - Downtown","Dubai - Business Bay",
        "Dubai - Jumeirah","Dubai - JBR","Dubai - Al Barsha","Dubai - Mirdif",
        "Dubai - Deira","Dubai - Bur Dubai","Dubai - Silicon Oasis",
        "Dubai - Academic City","Dubai - Arabian Ranches","Dubai - International City",
        "Dubai - Dubai Hills","Dubai - The Springs","Dubai - The Meadows",
        "Dubai - The Greens","Dubai - Palm Jumeirah","Dubai - Al Qusais",
        "Dubai - Al Nahda","Dubai - JVC","Dubai - Damac Hills",
        "Dubai - Discovery Gardens","Dubai - IMPZ","Dubai - Al Warqa",
        "Dubai - Nad Al Sheba",
        "Sharjah - Al Majaz","Sharjah - Al Nahda","Sharjah - Al Taawun",
        "Sharjah - Muwaileh","Sharjah - Al Khan","Sharjah - Al Yarmook",
        "Sharjah - Al Qasimia","Sharjah - Al Fisht","Sharjah - Al Nasserya",
        "Sharjah - Al Goaz","Sharjah - Al Jubail","Sharjah - Maysaloon",
        "Ajman - Al Rashidiya","Ajman - Al Nuaimiya","Ajman - Al Mowaihat",
        "Ajman - Al Rawda","Ajman - Al Jurf","Ajman - Al Hamidiya",
        "Ajman - Al Rumailah","Ajman - Al Bustan","Ajman - City Center",
        "RAK - Al Nakheel","RAK - Al Dhait","RAK - Julph",
        "RAK - Khuzam","RAK - Al Qusaidat","RAK - Seih Al Uraibi",
        "RAK - Al Rams","RAK - Al Mairid","RAK - Mina Al Arab",
        "RAK - Al Hamra Village","RAK - Marjan Island",
        "Fujairah - Al Faseel","Fujairah - Madhab","Fujairah - Dibba",
        "Fujairah - Sakamkam","Fujairah - Mirbah","Fujairah - Al Taween",
        "Fujairah - Kalba","Fujairah - Qidfa","Fujairah - Al Aqah",
        "UAQ - Al Salama","UAQ - Al Haditha","UAQ - Al Raas",
        "UAQ - Al Dar Al Baida","UAQ - Al Khor","UAQ - Al Ramlah",
        "UAQ - Al Maidan","UAQ - Emirates City",
    ]

    # =========================
    # Setup
    # =========================
    catalog = None
    if _db is not None:
        try:
            rows = _db.db_query(
                'SELECT device as "Device", description as "Description", unit_price as "UnitPrice", warranty as "Warranty", image_base64 as "ImageBase64", image_path as "ImagePath" FROM products ORDER BY id'
            )
            if rows:
                catalog = pd.DataFrame(rows)
        except Exception:
            catalog = None
    if catalog is None:
        try:
            catalog = pd.read_excel("data/products.xlsx")
        except Exception:
            st.error("❌ ERROR: Cannot load product catalog")
            return

    required_cols = ["Device", "Description", "UnitPrice", "Warranty"]
    for col in required_cols:
        if col not in catalog.columns:
            st.error(f"❌ Missing column: {col}")
            return

    # Records helpers (match invoice logic)
    def load_records():
        # Try DB first, fallback to Excel
        if _db is not None:
            try:
                rows = _db.db_query('SELECT base_id, date, type, number, amount, client_name, phone, location, note FROM records ORDER BY date')
                if rows:
                    df = pd.DataFrame(rows)
                    df.columns = [c.strip().lower() for c in df.columns]
                    return df
            except Exception:
                pass
        try:
            df = pd.read_excel("data/records.xlsx")
            df.columns = [c.strip().lower() for c in df.columns]
            return df
        except:
            return pd.DataFrame(columns=[
                "base_id","date","type","number","amount","client_name","phone","location","note"
            ])

    def save_record(rec: dict):
        # Try saving to DB, fallback to Excel
        if _db is not None:
            try:
                # Delete existing by type+number then insert
                if rec.get('type') and rec.get('number'):
                    try:
                        _db.db_execute('DELETE FROM records WHERE type = %s AND number = %s', (rec.get('type'), rec.get('number')))
                    except Exception:
                        pass
                _db.db_execute(
                    'INSERT INTO records(base_id, date, type, number, amount, client_name, phone, location, note) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)',
                    (rec.get('base_id'), rec.get('date'), rec.get('type'), rec.get('number'), rec.get('amount'), rec.get('client_name'), rec.get('phone'), rec.get('location'), rec.get('note'))
                )
                return
            except Exception:
                pass

        df = load_records()
        if not df.empty and {"type", "number"}.issubset(df.columns):
            df = df[~((df["type"] == rec.get("type")) & (df["number"] == rec.get("number")))]
        df = pd.concat([df, pd.DataFrame([rec])], ignore_index=True)
        if {"type", "number"}.issubset(df.columns):
            df = df.drop_duplicates(subset=["type", "number"], keep="last")
        df.to_excel("data/records.xlsx", index=False)

    # Customers helpers (auto add from quotation)
    def ensure_customers_file():
        os.makedirs("data", exist_ok=True)
        path = "data/customers.xlsx"
        if not os.path.exists(path):
            cols = [
                "client_name","phone","location","email","status",
                "notes","tags","next_follow_up","assigned_to","last_activity"
            ]
            pd.DataFrame(columns=cols).to_excel(path, index=False)

    def load_customers():
        # Try DB first then fallback to Excel
        if _db is not None:
            try:
                rows = _db.db_query('SELECT id, name, phone, email, address FROM customers ORDER BY id')
                if rows:
                    df = pd.DataFrame(rows)
                    df = df.rename(columns={'name':'client_name', 'address':'location'})
                    df.columns = [c.strip().lower() for c in df.columns]
                    return df
            except Exception:
                pass
        ensure_customers_file()
        try:
            df = pd.read_excel("data/customers.xlsx")
            df.columns = [c.strip().lower() for c in df.columns]
            return df
        except:
            return pd.DataFrame(columns=[
                "client_name","phone","location","email","status",
                "notes","tags","next_follow_up","assigned_to","last_activity"
            ])

    def save_customers(df: pd.DataFrame):
        os.makedirs("data", exist_ok=True)
        # Try DB sync (upsert) then write Excel to preserve app-specific fields
        if _db is not None:
            try:
                for _, row in df.iterrows():
                    name = str(row.get('client_name') or '')
                    phone = row.get('phone')
                    email = row.get('email')
                    address = row.get('location')
                    try:
                        existing = _db.db_query('SELECT id FROM customers WHERE name = %s AND phone = %s', (name, phone))
                    except Exception:
                        existing = []
                    if existing:
                        try:
                            _db.db_execute('UPDATE customers SET email = %s, address = %s WHERE id = %s', (email, address, existing[0].get('id')))
                        except Exception:
                            pass
                    else:
                        try:
                            _db.db_execute('INSERT INTO customers(name, phone, email, address) VALUES (%s,%s,%s,%s)', (name, phone, email, address))
                        except Exception:
                            pass
                df.to_excel("data/customers.xlsx", index=False)
                return
            except Exception:
                pass
        df.to_excel("data/customers.xlsx", index=False)

    def upsert_customer_from_quotation(name: str, phone: str, location: str):
        if not str(name).strip():
            return
        # Try DB upsert matching DB schema, otherwise fall back to Excel logic
        if _db is not None:
            try:
                # Simple phone/name lookup
                existing = []
                try:
                    existing = _db.db_query('SELECT id FROM customers WHERE name = %s OR phone = %s', (proper_case(name), phone))
                except Exception:
                    existing = []
                if existing:
                    try:
                        _db.db_execute('UPDATE customers SET phone = %s, address = %s WHERE id = %s', (phone, proper_case(location), existing[0].get('id')))
                    except Exception:
                        pass
                else:
                    try:
                        _db.db_execute('INSERT INTO customers(name, phone, email, address) VALUES (%s,%s,%s,%s)', (proper_case(name), phone, '', proper_case(location)))
                    except Exception:
                        pass
                return
            except Exception:
                pass

        # Fallback: original Excel behaviour
        ensure_customers_file()
        cdf = load_customers()
        key = str(name).strip().lower()
        exists = None
        if not cdf.empty and "client_name" in cdf.columns:
            m = cdf["client_name"].astype(str).str.strip().str.lower() == key
            if m.any():
                exists = cdf[m].index[0]
            else:
                # Try phone-based matching when names differ
                cdf_phone = cdf.get("phone", pd.Series([], dtype=str)).apply(lambda x: x if pd.isna(x) else str(x))
                try_phone = phone
                def norm(x):
                    digits = ''.join(filter(str.isdigit, str(x)))
                    if digits.startswith('971') and len(digits) >= 12 and digits[3] == '5':
                        return '0' + digits[3:12]
                    if len(digits) == 9 and digits.startswith('5'):
                        return '0' + digits
                    if len(digits) == 10 and digits.startswith('05'):
                        return digits
                    return digits[-10:]
                if try_phone:
                    m2 = cdf_phone.apply(norm) == norm(try_phone)
                    if m2.any():
                        try:
                            # Inline buttons: HTML download + Headless PDF (same line)
                            col_html, col_pdf = st.columns([1,1])
                            html_path = None

                            def _run_converter_background(html_path_str: str, status_path_str: str, vw: str | None = None, vh: str | None = None):
                                """Background worker to run the converter and write status JSON."""
                                status = {"status": "running", "started_at": time.time()}
                                try:
                                    Path(status_path_str).write_text(json.dumps(status), encoding='utf-8')
                                    script = Path(__file__).resolve().parents[1] / 'scripts' / 'convert_with_playwright.py'
                                    cmd = [sys.executable, str(script), html_path_str]
                                    if vw and vh:
                                        cmd.extend([str(int(vw)), str(int(vh))])
                                    proc = subprocess.run(cmd, check=True, capture_output=True, text=True)
                                    # On success, determine pdf path
                                    html_p = Path(html_path_str)
                                    pdf_p = html_p.with_suffix('.pdf')
                                    if pdf_p.exists():
                                        status.update({"status": "done", "pdf": str(pdf_p), "finished_at": time.time()})
                                    else:
                                        status.update({"status": "error", "error": "PDF not found after conversion", "finished_at": time.time()})
                                except Exception as exc:
                                    status.update({"status": "error", "error": str(exc), "finished_at": time.time()})
                                try:
                                    Path(status_path_str).write_text(json.dumps(status), encoding='utf-8')
                                except Exception:
                                    pass

                            if col_pdf.button('Export PDF (Headless)'):
                                # Ensure HTML is prepared
                                html_path = None
                                try:
                                    html_content = render_quotation_html({
                                        'company_name': load_settings().get('company_name', 'Newton Smart Home'),
                                        'quotation_number': quote_no,
                                        'quotation_date': datetime.today().strftime('%Y-%m-%d'),
                                        'valid_until': '',
                                        'status': 'Pending Approval',
                                        'client_name': client_name,
                                        'client_company': '',
                                        'client_address': client_location,
                                        'client_city': '',
                                        'client_trn': '',
                                        'project_title': '',
                                        'project_location': client_location,
                                        'project_scope': '',
                                        'project_notes': '',
                                        'items': st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else [],
                                        'subtotal': product_total,
                                        'Installation': float(st.session_state.get('install_cost_quo_value', 0.0) or 0.0),
                                        'vat_amount': 0,
                                        'total_amount': grand_total,
                                        'bank_name': load_settings().get('bank_name', ''),
                                        'bank_account': load_settings().get('bank_account', ''),
                                        'bank_iban': load_settings().get('bank_iban', ''),
                                        'bank_company': load_settings().get('company_name', 'Newton Smart Home'),
                                        'sig_name': load_settings().get('default_prepared_by', ''),
                                        'sig_role': load_settings().get('default_approved_by', ''),
                                    }, template_name="newton_quotation_A4.html")
                                    out_dir = Path('data') / 'exports'
                                    out_dir.mkdir(parents=True, exist_ok=True)
                                    html_path = out_dir / f"Quotation_{client_name}_{quote_no}.html"
                                    with open(html_path, 'w', encoding='utf-8') as fh:
                                        fh.write(html_content)
                                except Exception as _e:
                                    st.error(f"Failed to prepare HTML for headless export: {_e}")

                                # Start background thread that runs converter and writes a status file next to HTML
                                if html_path:
                                    status_path = html_path.with_suffix('.status.json')
                                    # Remove any previous status
                                    try:
                                        if status_path.exists():
                                            status_path.unlink()
                                    except Exception:
                                        pass
                                    vw = os.environ.get('PLAYWRIGHT_PDF_WIDTH')
                                    vh = os.environ.get('PLAYWRIGHT_PDF_HEIGHT')
                                    t = threading.Thread(target=_run_converter_background, args=(str(html_path), str(status_path), vw, vh), daemon=True)
                                    t.start()
                                    st.success('Headless conversion started in background. Click "Refresh status" to update.')

                            # Status area and refresh
                            try:
                                # Look for any recent status file for the current quotation
                                out_dir = Path('data') / 'exports'
                                expected_html = out_dir / f"Quotation_{client_name}_{quote_no}.html"
                                status_file = expected_html.with_suffix('.status.json')
                                if status_file.exists():
                                    try:
                                        s = json.loads(status_file.read_text(encoding='utf-8'))
                                    except Exception:
                                        s = {'status': 'unknown'}
                                    if s.get('status') == 'running':
                                        st.info('Conversion status: running')
                                        if st.button('Refresh status'):
                                            st.experimental_rerun()
                                    elif s.get('status') == 'done':
                                        pdf_path = Path(s.get('pdf')) if s.get('pdf') else expected_html.with_suffix('.pdf')
                                        if pdf_path.exists():
                                            with open(pdf_path, 'rb') as pf:
                                                pdf_bytes = pf.read()
                                            st.success('Conversion finished. Download below:')
                                            st.download_button('Download Headless PDF', pdf_bytes, file_name=pdf_path.name, mime='application/pdf')
                                        else:
                                            st.warning('Conversion reported done but PDF file missing.')
                                    elif s.get('status') == 'error':
                                        st.error(f"Conversion failed: {s.get('error')}")
                                        if st.button('Refresh status'):
                                            st.experimental_rerun()
                            except Exception:
                                pass
                        except Exception as e:
                            st.error(f'Export (headless) unavailable: {e}')

    st.session_state.num_entries = 1

    df = st.session_state.product_table.copy()

    if not df.empty:
        for idx, (i, row) in enumerate(df.iterrows()):
            cols = st.columns([4.5,0.7,1,1,0.7,0.7])

            with cols[0]:
                st.markdown(f"""
                    <div class='added-product-row'>
                        <span style="font-weight:bold;color:rgba(10,132,255,.65);">✓</span>
                        <span style="font-weight:600;color:#1f2937;">{row['Product / Device']}</span>
                    </div>
                """, unsafe_allow_html=True)

            with cols[1]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{int(row['Qty'])}</span></div>", unsafe_allow_html=True)

            with cols[2]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{row['Unit Price (AED)']:.2f}</span></div>", unsafe_allow_html=True)

            with cols[3]:
                st.markdown(
                    f"<div class='added-product-row'><span class='product-value'>AED {row['Line Total (AED)']:.2f}</span></div>",
                    unsafe_allow_html=True
                )

            with cols[4]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{int(row['Warranty (Years)'])} yr</span></div>", unsafe_allow_html=True)

            with cols[5]:
                if st.button("❌", key=f"del_q_{i}"):
                    st.session_state.product_table = st.session_state.product_table.drop(i).reset_index(drop=True)
                    st.session_state.product_table["Item No"] = range(1, len(st.session_state.product_table)+1)
                    st.rerun()

    for entry_idx in range(st.session_state.num_entries):
        cols = st.columns([4.5,0.7,1,1,0.7,0.7])

        with cols[0]:
            product = st.selectbox(
                "Product",
                catalog["Device"],
                key=f"prod_entry_{entry_idx}",
                label_visibility="collapsed"
            )
            row = catalog[catalog["Device"] == product].iloc[0]
            desc = row["Description"]

        key_qty = f"qty_val_{entry_idx}"
        key_price = f"price_val_{entry_idx}"
        key_war = f"war_val_{entry_idx}"
        if key_qty not in st.session_state:
            st.session_state[key_qty] = 1
        if key_price not in st.session_state:
            st.session_state[key_price] = float(row["UnitPrice"])
        if key_war not in st.session_state:
            st.session_state[key_war] = int(row["Warranty"])
        # Sync price and warranty when product changes
        last_key = f"last_prod_{entry_idx}"
        if st.session_state.get(last_key) != product:
            st.session_state[f"price_val_{entry_idx}"] = float(row["UnitPrice"])
            st.session_state[f"war_val_{entry_idx}"] = int(row["Warranty"])
            st.session_state[last_key] = product

        with cols[1]:
            st.number_input(
                "Qty",
                min_value=1,
                step=1,
                key=key_qty,
                label_visibility="collapsed"
            )

        with cols[2]:
            st.number_input(
                "Unit Price (AED)",
                min_value=0.0,
                step=10.0,
                key=key_price,
                label_visibility="collapsed"
            )

        qty = st.session_state[f"qty_val_{entry_idx}"]
        price = st.session_state[f"price_val_{entry_idx}"]
        line_price = qty * price

        with cols[3]:
            st.markdown(
                f"<div class='added-product-row'><span class='product-value'>AED {line_price:.2f}</span></div>",
                unsafe_allow_html=True
            )

        with cols[4]:
            st.number_input(
                "Warranty (Years)",
                min_value=0,
                step=1,
                key=key_war,
                label_visibility="collapsed"
            )

        warranty = st.session_state[f"war_val_{entry_idx}"]

        with cols[5]:
            if st.button("✅", key=f"add_row_{entry_idx}"):
                # attempt to attach image info from catalog (prefer Base64)
                image_val = None
                try:
                    if 'ImageBase64' in catalog.columns and not pd.isna(row.get('ImageBase64')):
                        raw_b64 = str(row.get('ImageBase64')).strip()
                        if raw_b64:
                            image_val = raw_b64 if raw_b64.startswith('data:') else f"data:image/png;base64,{raw_b64}"
                    if image_val is None and 'ImagePath' in catalog.columns and not pd.isna(row.get('ImagePath')):
                        image_val = str(row.get('ImagePath')).strip()
                except Exception:
                    image_val = None

                new_row = {
                    "Item No": len(st.session_state.product_table) + 1,
                    "Product / Device": product,
                    "Description": desc,
                    "Qty": qty,
                    "Unit Price (AED)": price,
                    "Line Total (AED)": line_price,
                    "Warranty (Years)": warranty,
                    # keep both raw columns for Word export and a normalized `image` for HTML rendering
                    "ImagePath": row.get('ImagePath') if 'ImagePath' in row.index else None,
                    "ImageBase64": row.get('ImageBase64') if 'ImageBase64' in row.index else None,
                    "image": image_val,
                }
                st.session_state.product_table = pd.concat(
                    [st.session_state.product_table, pd.DataFrame([new_row])],
                    ignore_index=True
                )
                st.rerun()

    st.markdown("---")

    product_total = st.session_state.product_table["Line Total (AED)"].sum() if not st.session_state.product_table.empty else 0

    # =========================
    # SUMMARY (match invoice)
    # =========================
    st.markdown("---")
    col_left, col_right = st.columns([1, 1])

    with col_left:
        st.markdown("<div class='section-title'>Project Costs</div>", unsafe_allow_html=True)

        # Pull persisted values (so left card reflects right inputs)
        installation_cost_val = st.session_state.get("install_cost_quo_value", 0.0)
        discount_value_val = st.session_state.get("disc_value_quo_value", 0.0)
        discount_percent_val = st.session_state.get("disc_percent_quo_value", 0.0)

        percent_value = (product_total + installation_cost_val) * (discount_percent_val / 100)
        total_discount = percent_value + discount_value_val
        grand_total = (product_total + installation_cost_val) - total_discount

        st.markdown(
            """
            <div style='background:#fff;border:1px solid rgba(0,0,0,.08);border-radius:12px;padding:16px;box-shadow:0 2px 6px rgba(0,0,0,.04);'>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Price (AED)</span>
                    <span style='font-weight:700;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Installation & Operation Devices</span>
                    <span style='font-weight:700;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Discount</span>
                    <span style='font-weight:700;color:#1d1d1f;'>-{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:15px 0;background:rgba(0,0,0,.02);margin-top:8px;border-radius:8px;padding-left:12px;padding-right:12px;'>
                    <span style='font-weight:700;font-size:16px;color:#1d1d1f;'>TOTAL AMOUNT</span>
                    <span style='font-weight:700;font-size:18px;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
            </div>
            """.format(product_total, installation_cost_val, total_discount, grand_total),
            unsafe_allow_html=True,
        )

    with col_right:
        st.markdown("<div class='section-title'>Installation & Discount</div>", unsafe_allow_html=True)

        installation_cost = st.number_input(
            "Installation & Operation Devices (AED)",
            min_value=0.0,
            step=50.0,
            key="install_cost_quo",
        )
        st.session_state["install_cost_quo_value"] = installation_cost

        st.markdown("<div style='margin-top:12px;'></div>", unsafe_allow_html=True)
        cD1, cD2 = st.columns(2)
        with cD1:
            discount_value = st.number_input("Discount Value (AED)", min_value=0.0, key="disc_value_quo")
            st.session_state["disc_value_quo_value"] = discount_value
        with cD2:
            discount_percent = st.number_input("Discount %", min_value=0.0, max_value=100.0, key="disc_percent_quo")
            st.session_state["disc_percent_quo_value"] = discount_percent

    # =========================
    # EXPORT HELPERS (on-click only)
    # =========================
    def generate_word_file(data: dict) -> BytesIO:
        doc = Document("data/quotation_template.docx")

        # قراءة أبعاد الصور من الإعدادات (سم)
        _s = load_settings()
        _wcm = float(_s.get("quote_product_image_width_cm", 3.49))
        _hcm = float(_s.get("quote_product_image_height_cm", 1.5))

        # خريطة أسماء المنتجات إلى صورة Base64 أو مسارها الأصلي (إن وُجدت)
        image_map = {}
        image_path_map = {}
        try:
            if 'ImageBase64' in catalog.columns:
                image_map = dict(zip(catalog['Device'].astype(str), catalog['ImageBase64']))
            if 'ImagePath' in catalog.columns:
                image_path_map = dict(zip(catalog['Device'].astype(str), catalog['ImagePath']))
        except Exception:
            image_map = {}
            image_path_map = {}

        def insert_image_in_cell(cell, b64_str: str, width_cm: float, height_cm: float, img_path: str = None):
            try:
                bio = None
                if img_path and os.path.exists(img_path):
                    with open(img_path, "rb") as f:
                        bio = BytesIO(f.read())
                elif b64_str and not pd.isna(b64_str):
                    img_bytes = base64.b64decode(b64_str)
                    bio = BytesIO(img_bytes)
                if bio is None:
                    return False
                # تفريغ محتوى الخلية ثم إدراج الصورة في فقرة محاذاة للوسط
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(bio, width=Cm(width_cm), height=Cm(height_cm))
                return True
            except Exception:
                return False

        def format_cell(cell, font_name, font_size, align):
            for paragraph in cell.paragraphs:
                paragraph.alignment = align
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    old = cell.text
                    new = old
                    for key, val in data.items():
                        new = new.replace(key, "" if val is None else str(val))
                    if old != new:
                        cell.text = new
                        format_cell(cell, "Times New Roman (Headings CS)", 10, WD_ALIGN_PARAGRAPH.LEFT)

        # Insert products from session state
        products = st.session_state.product_table.to_dict("records") if "product_table" in st.session_state else []

        target_table = None
        for table in doc.tables:
            try:
                if table.cell(0, 0).text.strip().lower() in ["item no", "item no."]:
                    target_table = table
                    break
            except Exception:
                continue
        if not target_table:
            raise Exception("❌ Product table not found")

        start_row = 1
        last_index = None
        for idx, row in enumerate(target_table.rows):
            if row.cells[0].text.strip().lower() == "last":
                last_index = idx
                break
        if last_index is None:
            raise Exception("❌ 'last' row missing in Word template")

        for i, product in enumerate(products):
            row_index = start_row + i
            if row_index >= last_index:
                break
            row = target_table.rows[row_index]
            row.cells[0].text = str(product.get("Item No", i + 1))
            # إدراج الصورة في عمود المنتج إن وُجدت، وإلا نكتب الاسم نصياً
            prod_name = str(product.get("Product / Device", ""))
            b64_img = image_map.get(prod_name)
            img_path = image_path_map.get(prod_name)
            placed = insert_image_in_cell(row.cells[1], b64_img, _wcm, _hcm, img_path)
            if not placed:
                row.cells[1].text = prod_name
            row.cells[2].text = str(product.get("Description", ""))
            row.cells[3].text = str(product.get("Qty", ""))
            row.cells[4].text = f"{float(product.get('Unit Price (AED)', 0)):,.2f}"
            row.cells[5].text = f"{float(product.get('Line Total (AED)', 0)):,.2f}"
            row.cells[6].text = str(product.get("Warranty (Years)", ""))
            for cell in row.cells:
                format_cell(cell, "Arial MT", 9, WD_ALIGN_PARAGRAPH.CENTER)

        delete_start = start_row + len(products)
        for j in range(last_index - 1, delete_start - 1, -1):
            row = target_table.rows[j]
            target_table._tbl.remove(row._tr)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def convert_to_pdf(word_buffer: BytesIO) -> bytes:
        # Deprecated: DOCX->PDF via ConvertAPI removed.
        # Instead render HTML template and convert to PDF directly.
        products = st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else []
        data = {
            'client_name': st.session_state.get('quo_client_name', ''),
            'client_location': st.session_state.get('quo_loc', ''),
            'quote_no': st.session_state.get('quo_no', ''),
        }
        html = render_quotation_html({
            'company_name': load_settings().get('company_name', 'Newton Smart Home'),
            'quotation_number': data.get('quote_no', ''),
            'quotation_date': datetime.today().strftime('%Y-%m-%d'),
            'valid_until': '',
            'status': 'Pending Approval',
            'client_name': data.get('client_name', ''),
            'client_company': '',
            'client_address': data.get('client_location', ''),
            'client_city': '',
            'client_trn': '',
            'project_title': '',
            'project_location': data.get('client_location', ''),
            'project_scope': '',
            'project_notes': '',
            'items': products,
            'subtotal': sum([float(p.get('Line Total (AED)', 0) or 0) for p in products]),
            'Installation': float(st.session_state.get('install_cost_quo_value', 0.0) or 0.0),
            'vat_amount': 0,
            'total_amount': sum([float(p.get('Line Total (AED)', 0) or 0) for p in products]),
            'bank_name': load_settings().get('bank_name', ''),
            'bank_account': load_settings().get('bank_account', ''),
            'bank_iban': load_settings().get('bank_iban', ''),
            'bank_company': load_settings().get('company_name', 'Newton Smart Home'),
            'sig_name': load_settings().get('default_prepared_by', ''),
            'sig_role': load_settings().get('default_approved_by', ''),
        }, template_name="newton_quotation_A4.html")
        return html_to_pdf(html)

    def _auto_download(data_bytes: bytes, filename: str, mime: str):
        b64 = base64.b64encode(data_bytes).decode('utf-8')
        # Visible fallback link (in case browser blocks auto-download)
        st.markdown(
            f"If the download doesn't start, click here: [Download {filename}](data:{mime};base64,{b64})",
            unsafe_allow_html=True,
        )
        # Build a compact JS snippet for auto-download. Escape backslashes carefully.
        js = (
            "<script>"
            "(function(){"
            f"const b64='{b64}';"
            f"const mime='{mime}';"
            "const byteChars=atob(b64);"
            "const byteNumbers=new Array(byteChars.length);"
            "for(let i=0;i<byteChars.length;i++)byteNumbers[i]=byteChars.charCodeAt(i);"
            "const byteArray=new Uint8Array(byteNumbers);"
            "const blob=new Blob([byteArray],{type:mime});"
            "const url=URL.createObjectURL(blob);"
            "const a=document.createElement('a');a.href=url;"
            f"a.download='{filename}';"
            "document.body.appendChild(a);a.click();"
            "setTimeout(function(){URL.revokeObjectURL(url);a.remove();},1000);"
            "})();</script>"
        )
        st_html(js, height=0)

    def _save_export_locally(data_bytes: bytes, filename: str) -> str:
        out_dir = Path('data') / 'exports'
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / filename
        with open(out_path, 'wb') as f:
            f.write(data_bytes)
        return str(out_path)

    st.markdown("---")
    st.markdown('<div class="section-title">Export Quotation</div>', unsafe_allow_html=True)

    # Button colors (blue for Word, red for PDF)
    st.markdown(
        """
        <style>
        div.stButton>button[k="word_action"]{
            background:linear-gradient(145deg,#0a84ff 0%,#1b6cff 100%)!important;color:#fff!important;
            border:1px solid rgba(10,132,255,.35)!important;border-radius:12px!important;
            padding:8px 16px!important;font-weight:700!important;
        }
        div.stButton>button[k="pdf_action"]{
            background:linear-gradient(145deg,#ff3b30 0%,#d70015 100%)!important;color:#fff!important;
            border:1px solid rgba(255,59,48,.35)!important;border-radius:12px!important;
            padding:8px 16px!important;font-weight:700!important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Recalculate using the Installation & Discount values (to mirror invoice)
    product_total = st.session_state.product_table["Line Total (AED)"].sum() if not st.session_state.product_table.empty else 0.0
    installation_cost_val = st.session_state.get("install_cost_quo_value", 0.0)
    discount_value_val = st.session_state.get("disc_value_quo_value", 0.0)
    discount_percent_val = st.session_state.get("disc_percent_quo_value", 0.0)
    percent_value = (product_total + installation_cost_val) * (discount_percent_val / 100)
    total_discount = percent_value + discount_value_val
    grand_total = (product_total + installation_cost_val) - total_discount

    data_to_fill = {
        "{{client_name}}": client_name,
        "{{quote_no}}": quote_no,
        "{{client_location}}": client_location,
        "{{prepared_by}}": prepared_by,
        "{{client_phone}}": client_phone or "N/A",
        "{{approved_by}}": approved_by,
        "{{client_email}}": "N/A",
        # Quotation template keys
        "{{total1}}": f"{product_total:,.2f}",
        "{{installation_cost}}": f"{installation_cost_val:,.2f}",
        "{{Price}}": f"{product_total:,.2f}",
        "{{Total}}": f"{grand_total:,.2f}",
        # Extra keys (no-op if not present in template)
        "{{discount_value}}": f"{discount_value_val:,.2f}",
        "{{discount_percent}}": f"{discount_percent_val:,.0f}",
        "{{total_discount}}": f"{total_discount:,.2f}",
        "{{grand_total}}": f"{grand_total:,.2f}",
    }

    # Always show the two action buttons side-by-side
    b1, b2 = st.columns(2)

    # Simple, invoice-style: pre-render a download_button for Word
    with b1:
        try:
            word_ready = generate_word_file(data_to_fill)
            clicked_word = st.download_button(
                label="Download Word",
                data=word_ready.getvalue(),
                file_name=f"Quotation_{client_name}_{quote_no}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_word_{quote_no}"
            )
            if clicked_word:
                # Save record after user downloads (same behavior as invoice)
                today_id = datetime.today().strftime('%Y%m%d')
                existing = load_records()
                if not existing.empty and "base_id" in existing.columns:
                    same_day = existing[existing.get("base_id", "").astype(str).str.contains(today_id, na=False)]
                    seq = len(same_day) + 1
                else:
                    seq = 1
                base_id = f"{today_id}-{str(seq).zfill(3)}"
                save_record({
                    "base_id": base_id,
                    "date": datetime.today().strftime('%Y-%m-%d'),
                    "type": "q",
                    "number": quote_no,
                    "amount": grand_total,
                    "client_name": client_name,
                    "phone": phone_raw,
                    "location": client_location,
                    "note": ""
                })
                upsert_customer_from_quotation(client_name, phone_raw, client_location)
                # Attempt to persist quotation and items to DB (non-intrusive)
                if _db is not None:
                    try:
                        # Ensure customer exists (upsert already attempted above)
                        cust = None
                        try:
                            cust_rows = _db.db_query('SELECT id FROM customers WHERE name = %s AND phone = %s', (proper_case(client_name), phone_raw))
                            cust = cust_rows[0].get('id') if cust_rows else None
                        except Exception:
                            cust = None

                        # Insert quotation and items
                        try:
                            qrow = _db.db_execute(
                                'INSERT INTO quotations(quote_number, customer_id, subtotal, installation_fee, total_amount, status, notes) VALUES (%s,%s,%s,%s,%s,%s,%s) RETURNING id',
                                (quote_no, cust, product_total, installation_cost_val, grand_total, 'pending', ''),
                                returning=True,
                            )
                            quotation_id = qrow.get('id') if qrow else None
                        except Exception:
                            quotation_id = None

                        products = st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else []
                        if quotation_id is not None and products:
                            for p in products:
                                try:
                                    prod_name = p.get('Product / Device')
                                    prod_rows = _db.db_query('SELECT id FROM products WHERE lower(device) = lower(%s) LIMIT 1', (prod_name,))
                                    prod_id = prod_rows[0].get('id') if prod_rows else None
                                except Exception:
                                    prod_id = None
                                try:
                                    _db.db_execute(
                                        'INSERT INTO quotation_items(quotation_id, product_id, description, quantity, unit_price, line_total, warranty) VALUES (%s,%s,%s,%s,%s,%s,%s)',
                                        (
                                            quotation_id,
                                            prod_id,
                                            p.get('Description'),
                                            p.get('Qty') or 0,
                                            p.get('Unit Price (AED)') or 0,
                                            p.get('Line Total (AED)') or 0,
                                            str(p.get('Warranty (Years)') or ''),
                                        ),
                                    )
                                except Exception:
                                    # Don't block on item-level failures
                                    pass
                        # Optionally track the export
                        try:
                            _db.db_execute('INSERT INTO exports(quotation_id, export_type, file_path, metadata) VALUES (%s,%s,%s,%s)', (quotation_id, 'word', '', None))
                        except Exception:
                            pass
                    except Exception:
                        # If any DB error occurs, fall back silently to Excel behaviour
                        pass

                # Log quotation creation
                user = st.session_state.get("user", {})
                log_event(user.get("name", "Unknown"), "Quotation", "quotation_created", 
                         f"Client: {client_name}, Amount: {grand_total}")
                st.success(f"✅ Saved quotation to records with base {base_id}")
        except Exception as e:
            st.error(f"❌ Unable to prepare Word file: {e}")

    # PDF: keep a click-to-generate then download button for reliability
    with b2:
        try:
            # Pre-generate PDF bytes for single-click download (may take a moment)
            word_for_pdf = generate_word_file(data_to_fill)
            pdf_ready = convert_to_pdf(word_for_pdf)
            clicked_pdf = st.download_button(
                label="Download PDF",
                data=pdf_ready,
                file_name=f"Quotation_{client_name}_{quote_no}.pdf",
                mime="application/pdf",
                key=f"dl_pdf_{quote_no}"
            )
            if clicked_pdf:
                today_id = datetime.today().strftime('%Y%m%d')
                existing = load_records()
                if not existing.empty and "base_id" in existing.columns:
                    same_day = existing[existing.get("base_id", "").astype(str).str.contains(today_id, na=False)]
                    seq = len(same_day) + 1
                else:
                    seq = 1
                base_id = f"{today_id}-{str(seq).zfill(3)}"
                save_record({
                    "base_id": base_id,
                    "date": datetime.today().strftime('%Y-%m-%d'),
                    "type": "q",
                    "number": quote_no,
                    "amount": grand_total,
                    "client_name": client_name,
                    "phone": phone_raw,
                    "location": client_location,
                    "note": "PDF"
                })
                upsert_customer_from_quotation(client_name, phone_raw, client_location)
                # Attempt DB persistence of quotation and items (non-intrusive)
                if _db is not None:
                    try:
                        cust = None
                        try:
                            cust_rows = _db.db_query('SELECT id FROM customers WHERE name = %s AND phone = %s', (proper_case(client_name), phone_raw))
                            cust = cust_rows[0].get('id') if cust_rows else None
                        except Exception:
                            cust = None

                        qrow = None
                        try:
                            qrow = _db.db_execute(
                                'INSERT INTO quotations(quote_number, customer_id, subtotal, installation_fee, total_amount, status, notes) VALUES (%s,%s,%s,%s,%s,%s,%s) RETURNING id',
                                (quote_no, cust, product_total, installation_cost_val, grand_total, 'pending', ''),
                                returning=True,
                            )
                        except Exception:
                            qrow = None
                        quotation_id = qrow.get('id') if qrow else None
                        products = st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else []
                        if quotation_id is not None and products:
                            for p in products:
                                try:
                                    prod_name = p.get('Product / Device')
                                    prod_rows = _db.db_query('SELECT id FROM products WHERE lower(device) = lower(%s) LIMIT 1', (prod_name,))
                                    prod_id = prod_rows[0].get('id') if prod_rows else None
                                except Exception:
                                    prod_id = None
                                try:
                                    _db.db_execute(
                                        'INSERT INTO quotation_items(quotation_id, product_id, description, quantity, unit_price, line_total, warranty) VALUES (%s,%s,%s,%s,%s,%s,%s)',
                                        (
                                            quotation_id,
                                            prod_id,
                                            p.get('Description'),
                                            p.get('Qty') or 0,
                                            p.get('Unit Price (AED)') or 0,
                                            p.get('Line Total (AED)') or 0,
                                            str(p.get('Warranty (Years)') or ''),
                                        ),
                                    )
                                except Exception:
                                    pass
                        try:
                            _db.db_execute('INSERT INTO exports(quotation_id, export_type, file_path, metadata) VALUES (%s,%s,%s,%s)', (quotation_id, 'pdf', '', None))
                        except Exception:
                            pass
                    except Exception:
                        pass

                st.success(f"✅ Saved PDF quotation with base {base_id}")
        except Exception as e:
            st.error(f"❌ Unable to prepare PDF: {e}")

    # Headless export using Playwright (Chromium) - separate button
    try:
        # Inline buttons: HTML download + Headless PDF (same line)
        col_html, col_pdf = st.columns([1,1])
        html_path = None
        if col_pdf.button('Export PDF (Headless)'):
            # Ensure HTML is prepared
            html_path = None
            try:
                html_content = render_quotation_html({
                    'company_name': load_settings().get('company_name', 'Newton Smart Home'),
                    'quotation_number': quote_no,
                    'quotation_date': datetime.today().strftime('%Y-%m-%d'),
                    'valid_until': '',
                    'status': 'Pending Approval',
                    'client_name': client_name,
                    'client_company': '',
                    'client_address': client_location,
                    'client_city': '',
                    'client_trn': '',
                    'project_title': '',
                    'project_location': client_location,
                    'project_scope': '',
                    'project_notes': '',
                    'items': st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else [],
                    'subtotal': product_total,
                    'Installation': float(st.session_state.get('install_cost_quo_value', 0.0) or 0.0),
                    'vat_amount': 0,
                    'total_amount': grand_total,
                    'bank_name': load_settings().get('bank_name', ''),
                    'bank_account': load_settings().get('bank_account', ''),
                    'bank_iban': load_settings().get('bank_iban', ''),
                    'bank_company': load_settings().get('company_name', 'Newton Smart Home'),
                    'sig_name': load_settings().get('default_prepared_by', ''),
                    'sig_role': load_settings().get('default_approved_by', ''),
                }, template_name="newton_quotation_A4.html")
                out_dir = Path('data') / 'exports'
                out_dir.mkdir(parents=True, exist_ok=True)
                html_path = out_dir / f"Quotation_{client_name}_{quote_no}.html"
                with open(html_path, 'w', encoding='utf-8') as fh:
                    fh.write(html_content)
            except Exception as _e:
                st.error(f"Failed to prepare HTML for headless export: {_e}")

            # Call the playwright converter script which will produce a PDF in data/exports
            if html_path and html_path.exists():
                import subprocess
                script = Path(__file__).resolve().parents[1] / 'scripts' / 'convert_with_playwright.py'
                try:
                    # Pass the html_path as argument to the converter
                    cmd = [sys.executable, str(script), str(html_path)]
                    proc = subprocess.run(cmd, check=True, capture_output=True, text=True)
                    st.success('Headless export completed')
                except Exception as e:
                    st.error(f'Headless export failed: {e} - {getattr(e, "stdout", "")}')
                    proc = None

                # Try to locate output PDF (consistent with converter naming)
                # Converter names outputs after the input file stem
                pdf_path = out_dir / f"{html_path.stem}.pdf"
                if pdf_path.exists():
                    with open(pdf_path, 'rb') as pf:
                        pdf_bytes = pf.read()
                        try:
                            col_pdf.download_button('Download Headless PDF', pdf_bytes, file_name=pdf_path.name, mime='application/pdf')
                        except Exception:
                            # If download_button fails for any reason, show a minimal info message
                            st.info('PDF ready for download.')
                else:
                    st.error('Headless export completed but PDF not found')
    except Exception as e:
        st.error(f'Export (headless) unavailable: {e}')

    # Provide Download HTML button (separate row)
    try:
        html_content = render_quotation_html({
            'company_name': load_settings().get('company_name', 'Newton Smart Home'),
            'quotation_number': quote_no,
            'quotation_date': datetime.today().strftime('%Y-%m-%d'),
            'valid_until': '',
            'status': 'Pending Approval',
            'client_name': client_name,
            'client_company': '',
            'client_address': client_location,
            'client_city': '',
            'client_trn': '',
            'project_title': '',
            'project_location': client_location,
            'project_scope': '',
            'project_notes': '',
            'items': st.session_state.product_table.to_dict('records') if 'product_table' in st.session_state else [],
            'subtotal': product_total,
            'Installation': float(st.session_state.get('install_cost_quo_value', 0.0) or 0.0),
            'vat_amount': 0,
            'total_amount': grand_total,
            'bank_name': load_settings().get('bank_name', ''),
            'bank_account': load_settings().get('bank_account', ''),
            'bank_iban': load_settings().get('bank_iban', ''),
            'bank_company': load_settings().get('company_name', 'Newton Smart Home'),
            'sig_name': load_settings().get('default_prepared_by', ''),
            'sig_role': load_settings().get('default_approved_by', ''),
        }, template_name="newton_quotation_A4.html")
        # Render HTML now and show inline Download HTML button
        if col_html.button('Download HTML'):
            col_html.download_button('Download Quotation (HTML)', html_content, file_name=f"Quotation_{client_name}_{quote_no}.html", mime='text/html')
    except Exception as e:
        st.error(f"❌ Unable to prepare HTML: {e}")


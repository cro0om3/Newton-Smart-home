import os
import base64
from io import BytesIO
from datetime import datetime
from pathlib import Path
from string import Template

import pandas as pd
import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.settings import load_settings


# ==========================================
# DATA SETUP
# ==========================================
def ensure_product_file() -> Path:
    os.makedirs("data", exist_ok=True)
    products_path = Path("data/products.xlsx")
    if not products_path.exists():
        df = pd.DataFrame(
            columns=["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]
        )
        df.to_excel(products_path, index=False)
    return products_path


def proper_case(text):
    if text is None:
        return ""
    try:
        return str(text).strip().title()
    except Exception:
        return str(text)


def load_products() -> pd.DataFrame:
    ensure_product_file()
    try:
        df = pd.read_excel("data/products.xlsx")
        for col in ["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]:
            if col not in df.columns:
                df[col] = None
        return df[["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]]
    except Exception:
        return pd.DataFrame(
            columns=["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]
        )


def save_products(df: pd.DataFrame):
    os.makedirs("data", exist_ok=True)
    for col in ["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]:
        if col not in df.columns:
            df[col] = None
    df = df[["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]]
    df.to_excel("data/products.xlsx", index=False)


# ==========================================
# IMAGE HELPERS
# ==========================================
def image_to_base64(uploaded_file, target_size=None, mode="contain"):
    """
    Convert an uploaded image file to base64 with optional resize/crop.
    Uses contain mode by default to fit inside the box without cropping or upscaling.
    Flattens to a white background and saves as JPEG to avoid clipping and keep size small.
    """
    try:
        raw = Image.open(uploaded_file)

        # Always flatten on white to avoid dark/transparent backgrounds
        if raw.mode in ("RGBA", "LA"):
            base = Image.new("RGBA", raw.size, (255, 255, 255, 255))
            base.paste(raw, mask=raw.split()[-1])
            raw = base.convert("RGB")
        else:
            raw = raw.convert("RGB")

        if target_size is None:
            settings = load_settings()
            tw = int(settings.get("ui_product_image_width_px", 350))
            th = int(settings.get("ui_product_image_height_px", 195))
        else:
            tw, th = target_size

        # Clamp processing size to fixed display size
        tw = min(tw, 350)
        th = min(th, 195)

        img = raw
        img_w, img_h = img.size
        img_ratio = img_w / img_h
        target_ratio = tw / th if th else 1

        if mode == "cover":
            # Fill the box, cropping excess; allows upscaling if needed.
            if img_ratio > target_ratio:
                new_h = th
                new_w = max(1, int(new_h * img_ratio))
                img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                left = max(0, (new_w - tw) // 2)
                img = img.crop((left, 0, left + tw, th))
            else:
                new_w = tw
                new_h = max(1, int(new_w / img_ratio))
                img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                top = max(0, (new_h - th) // 2)
                img = img.crop((0, top, tw, top + th))
        else:  # contain (fits inside box, no upscale)
            img.thumbnail((tw, th), Image.Resampling.LANCZOS)
            canvas = Image.new("RGB", (tw, th), (255, 255, 255))
            offset = ((tw - img.width) // 2, (th - img.height) // 2)
            canvas.paste(img, offset)
            img = canvas

        # Compress JPEG on a quality ladder to keep under Excel cell limits
        data = None
        for q in (80, 70, 60, 50):
            buffered = BytesIO()
            img.convert("RGB").save(buffered, format="JPEG", quality=q, optimize=True)
            data = buffered.getvalue()
            # Stop once size is reasonable for Excel cells (~24KB raw bytes -> ~32KB base64)
            if len(data) <= 24000:
                break
        return base64.b64encode(data).decode()
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None


def save_original_image(uploaded_file, device_name: str) -> str:
    """
    Save the uploaded image as PNG (full quality) for use in Word exports.
    Returns the saved path, or None on failure.
    """
    try:
        os.makedirs("data/product_images", exist_ok=True)
        safe_name = "".join(c for c in device_name if c.isalnum() or c in (" ", "_", "-")).strip() or "product"
        file_name = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        out_path = Path("data/product_images") / file_name
        img = Image.open(uploaded_file)
        if img.mode in ("RGBA", "LA"):
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        img.save(out_path, format="PNG", optimize=True)
        return str(out_path)
    except Exception:
        return None


def insert_product_card(doc: Document, row: pd.Series, width_cm: float, height_cm: float, card_index: int):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # Image cell
    img_cell = table.cell(0, 0)
    img_cell.merge(table.cell(1, 0))
    b64 = row.get("ImageBase64")
    if b64 and pd.notna(b64):
        try:
            img_bytes = base64.b64decode(b64)
            bio = BytesIO(img_bytes)
            img_cell.text = ""
            p = img_cell.paragraphs[0] if img_cell.paragraphs else img_cell.add_paragraph("")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(bio, width=Cm(width_cm), height=Cm(height_cm))
        except Exception:
            img_cell.text = "No Image"
    else:
        img_cell.text = "No Image"

    # Text cell
    text_cell = table.cell(0, 1)
    text_cell.paragraphs[0].text = ""
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = p.add_run(str(row.get("Device", "")))
    name_run.bold = True
    desc_p = text_cell.add_paragraph(str(row.get("Description", "")))
    price_p = text_cell.add_paragraph(f"{row.get('UnitPrice', '')} AED")
    price_p.paragraph_format.space_after = Pt(0)
    warr_p = text_cell.add_paragraph(f"Warranty: {row.get('Warranty', '')}")
    warr_p.paragraph_format.space_after = Pt(0)

    # Spacing between cards and page break every 4
    doc.add_paragraph("")
    if (card_index + 1) % 4 == 0:
        doc.add_page_break()


def build_word_cards_document(products_df: pd.DataFrame) -> BytesIO:
    doc = Document("data/catalog_template.docx")
    width_cm = float(load_settings().get("quote_product_image_width_cm", 3.49))
    height_cm = float(load_settings().get("quote_product_image_height_cm", 1.5))

    for idx, (_, row) in enumerate(products_df.iterrows()):
        insert_product_card(doc, row, width_cm, height_cm, idx)

    return save_docx_to_buffer(doc)


def save_docx_to_buffer(doc: Document) -> BytesIO:
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def base64_to_image_html(base64_str, width=None, height=None):
    if width is None or height is None:
        s = load_settings()
        width = int(s.get("ui_product_image_width_px", 350))
        height = int(s.get("ui_product_image_height_px", 195))
    if base64_str and pd.notna(base64_str):
        return f'<img src="data:image/jpeg;base64,{base64_str}" class="product-img">'
    return (
        '<div class="product-img" style="display:flex; align-items:center; justify-content:center; color:#999; font-size:14px;">No Image</div>'
    )


# ==========================================
# PAGE
# ==========================================
def products_app():
    ensure_product_file()

    settings = load_settings()
    # Keep processing size from settings; display size is fixed downscaled
    wpx = int(settings.get("ui_product_image_width_px", 350))
    hpx = int(settings.get("ui_product_image_height_px", 195))
    # Fixed display size per request
    display_w = 350
    display_h = 195

    st.markdown("<div class='section-title'>Products</div>", unsafe_allow_html=True)

    css_template = Template(
        """
        <style>
        .products-header{
            display:flex; gap:1rem; padding:8px 0 12px; background:transparent;
            font-size:11px; font-weight:600; letter-spacing:.06em; text-transform:uppercase;
            margin-bottom:10px; align-items:center;
        }
        .products-header span{text-align:center;}
        .products-header span:nth-child(1){flex:2.6; text-align:center;}
        .products-header span:nth-child(2){flex:2; text-align:left;}
        .products-header span:nth-child(3){flex:3; text-align:left;}
        .products-header span:nth-child(4){flex:1; text-align:center;}
        .products-header span:nth-child(5){flex:1; text-align:center;}
        .products-header span:nth-child(6){flex:1; text-align:center;}

        .added-product-row{ padding:8px 12px; border-radius:12px; box-shadow:0 2px 6px rgba(0,0,0,.05); }
        .product-value{ font-weight:600; }
        .product-image-cell{ 
            display:flex; 
            align-items:center; 
            justify-content:center; 
            width:${dw}px;
            height:${dh}px;
            overflow: hidden;
            background: transparent;
            border:1px solid rgba(0,0,0,0.05);
            border-radius:14px;
            padding:6px;
        }
        
        .product-img {
            width:${dw}px;
            height:${dh}px;
            object-fit: contain;
            background: transparent;
            border:1px solid rgba(0,0,0,0.03);
            border-radius: 10px;
            padding: 0;
            display: block;
            margin: 0 auto;
        }
        </style>
        """
    )
    st.markdown(css_template.substitute(dw=display_w, dh=display_h), unsafe_allow_html=True)

    # ---------------- ADD NEW PRODUCT (TOP) ----------------
    df = load_products()
    st.markdown("<div class='section-title'>Add New Product</div>", unsafe_allow_html=True)
    with st.expander("Add product", expanded=False):
        img_col1, img_col2 = st.columns([1, 3])
        with img_col1:
            st.markdown("**Product Image**")
            _s = load_settings()
            _wpx = int(_s.get("ui_product_image_width_px", 350))
            _hpx = int(_s.get("ui_product_image_height_px", 195))
            uploaded_image = st.file_uploader(
                f"Upload Image ({_wpx}x{_hpx}px)",
                type=["jpg", "jpeg", "png"],
                key="_a_img",
                label_visibility="collapsed",
            )
            if uploaded_image:
                temp_b64 = image_to_base64(uploaded_image)
                if temp_b64:
                    st.markdown(base64_to_image_html(temp_b64), unsafe_allow_html=True)
                uploaded_image.seek(0)
            else:
                st.markdown(base64_to_image_html(None), unsafe_allow_html=True)

        with img_col2:
            ar1, ar2 = st.columns(2)
            with ar1:
                a_device = st.text_input("Device", value="", key="_a_dev")
                a_price = st.number_input(
                    "UnitPrice", min_value=0.0, step=10.0, value=0.0, key="_a_price"
                )
            with ar2:
                a_warranty = st.number_input(
                    "Warranty", min_value=0, step=1, value=0, key="_a_war"
                )
                a_desc = st.text_area("Description", value="", height=90, key="_a_desc")

            ac1, ac2, _ = st.columns([1, 1, 2])
            with ac1:
                if st.button("Add Product"):
                    cand = proper_case(a_device)
                    if not cand.strip():
                        st.warning("Device is required.")
                    elif (
                        not df[
                            df["Device"]
                            .astype(str)
                            .str.strip()
                            .str.lower()
                            == cand.strip().lower()
                        ].empty
                    ):
                        st.warning("Device must be unique.")
                    else:
                        img_b64 = image_to_base64(uploaded_image) if uploaded_image else None
                        if uploaded_image:
                            uploaded_image.seek(0)
                        img_path = save_original_image(uploaded_image, cand) if uploaded_image else None
                        new_row = {
                            "Device": cand,
                            "Description": a_desc,
                            "UnitPrice": a_price,
                            "Warranty": a_warranty,
                            "ImageBase64": img_b64,
                            "ImagePath": img_path,
                        }
                        new_df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
                        save_products(new_df)
                        st.success("Product added")
                        st.rerun()
            with ac2:
                if st.button("Reset Form"):
                    st.session_state["_a_dev"] = ""
                    st.session_state["_a_price"] = 0.0
                    st.session_state["_a_war"] = 0
                    st.session_state["_a_desc"] = ""
                    st.session_state.pop("_a_img", None)
                    st.rerun()

    st.markdown("---")

    # ---------------- SEARCH & FILTERS ----------------
    st.markdown(
        "<div class='section-title'>Quick Search</div>",
        unsafe_allow_html=True,
    )
    s1, s2 = st.columns([2, 1])
    with s1:
        q_text = st.text_input(
            "Search by name or description", placeholder="Type to filter products..."
        )
    with s2:
        only_with_images = st.checkbox("Show items with images only", value=False)

    fdf = df.copy()
    if q_text:
        ql = q_text.strip().lower()
        mask_name = fdf["Device"].astype(str).str.lower().str.contains(ql, na=False)
        mask_desc = fdf["Description"].astype(str).str.lower().str.contains(ql, na=False)
        fdf = fdf[mask_name | mask_desc]
    if only_with_images:
        fdf = fdf[fdf["ImageBase64"].notna() & (fdf["ImageBase64"].astype(str) != "")]

    # ---------------- TABLE ----------------
    st.markdown("<div class='section-title'>Catalog</div>", unsafe_allow_html=True)

    st.markdown(
        """
        <div class="product-header products-header">
            <span>Image</span>
            <span>Device</span>
            <span>Description</span>
            <span>Unit Price (AED)</span>
            <span>Warranty (Years)</span>
            <span>Actions</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if fdf.empty:
        st.info("No products found. Add your first product above.")
    else:
        for display_idx, (original_idx, row) in enumerate(fdf.iterrows()):
            is_editing = st.session_state.get("_prod_edit_idx") == original_idx
            dcol = st.columns([2.6, 2, 3, 1, 1, 1])

            if is_editing:
                with dcol[0]:
                    img_upload = st.file_uploader(
                        "Image",
                        type=["jpg", "jpeg", "png"],
                        key=f"edit_img_{display_idx}",
                        label_visibility="collapsed",
                    )
                    if img_upload:
                        img_b64 = image_to_base64(img_upload)
                        if img_b64:
                            st.markdown(
                                base64_to_image_html(img_b64), unsafe_allow_html=True
                            )
                        img_upload.seek(0)
                    else:
                        st.markdown(
                            base64_to_image_html(row.get("ImageBase64")),
                            unsafe_allow_html=True,
                        )

                with dcol[1]:
                    edit_device = st.text_input(
                        "Device",
                        value=row["Device"],
                        key=f"edit_dev_{display_idx}",
                        label_visibility="collapsed",
                    )
                with dcol[2]:
                    edit_desc = st.text_input(
                        "Description",
                        value=str(row["Description"]),
                        key=f"edit_desc_{display_idx}",
                        label_visibility="collapsed",
                    )
                with dcol[3]:
                    try:
                        default_price = float(row["UnitPrice"])
                    except Exception:
                        default_price = 0.0
                    edit_price = st.number_input(
                        "Price",
                        value=default_price,
                        min_value=0.0,
                        step=10.0,
                        key=f"edit_price_{display_idx}",
                        label_visibility="collapsed",
                    )
                with dcol[4]:
                    try:
                        default_warranty = int(row["Warranty"])
                    except Exception:
                        default_warranty = 0
                    edit_warranty = st.number_input(
                        "Warranty",
                        value=default_warranty,
                        min_value=0,
                        step=1,
                        key=f"edit_war_{display_idx}",
                        label_visibility="collapsed",
                    )
                with dcol[5]:
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Save", key=f"save_{display_idx}", help="Save changes"):
                            cand = proper_case(edit_device)
                            if not cand.strip():
                                st.warning("Device name is required.")
                            else:
                                dup = df[
                                    (df.index != original_idx)
                                    & (
                                        df["Device"]
                                        .astype(str)
                                        .str.strip()
                                        .str.lower()
                                        == cand.strip().lower()
                                    )
                                ]
                                if not dup.empty:
                                    st.warning("Device name must be unique.")
                                else:
                                    new_img_b64 = row.get("ImageBase64")
                                    new_img_path = row.get("ImagePath")
                                    if img_upload:
                                        new_img_b64 = image_to_base64(img_upload)
                                        img_upload.seek(0)
                                        new_img_path = save_original_image(img_upload, proper_case(edit_device))

                                    df.loc[
                                        original_idx,
                                        [
                                            "Device",
                                            "Description",
                                            "UnitPrice",
                                            "Warranty",
                                            "ImageBase64",
                                            "ImagePath",
                                        ],
                                    ] = [
                                        proper_case(edit_device),
                                        edit_desc,
                                        edit_price,
                                        edit_warranty,
                                        new_img_b64,
                                        new_img_path,
                                    ]
                                    save_products(df)
                                    st.session_state.pop("_prod_edit_idx", None)
                                    st.success("Product updated")
                                    st.rerun()
                    with c2:
                        if st.button("Cancel", key=f"cancel_{display_idx}", help="Cancel"):
                            st.session_state.pop("_prod_edit_idx", None)
                            st.rerun()
            else:
                with dcol[0]:
                    st.markdown(
                        f'<div class="product-image-cell">{base64_to_image_html(row.get("ImageBase64"))}</div>',
                        unsafe_allow_html=True,
                    )
                with dcol[1]:
                    st.markdown(
                        f"<div class='added-product-row' style='padding:8px 12px;'><span class='product-value' style='font-size:13px'>{row['Device']}</span></div>",
                        unsafe_allow_html=True,
                    )
                with dcol[2]:
                    st.markdown(
                        f"<div class='added-product-row' style='padding:8px 12px;font-size:13px'>{row['Description']}</div>",
                        unsafe_allow_html=True,
                    )
                with dcol[3]:
                    try:
                        price_txt = f"{float(row['UnitPrice']):,.2f}"
                    except Exception:
                        price_txt = str(row["UnitPrice"])
                    st.markdown(
                        f"<div class='added-product-row' style='padding:8px 12px;'><span class='product-value' style='font-size:13px'>{price_txt}</span></div>",
                        unsafe_allow_html=True,
                    )
                with dcol[4]:
                    try:
                        war_txt = f"{int(row['Warranty'])}"
                    except Exception:
                        war_txt = str(row["Warranty"])
                    st.markdown(
                        f"<div class='added-product-row' style='padding:8px 12px;'><span class='product-value' style='font-size:13px'>{war_txt}</span></div>",
                        unsafe_allow_html=True,
                    )
                with dcol[5]:
                    c1, c2 = st.columns(2)
                    with c1:
                        if st.button("Edit", key=f"edit_{display_idx}"):
                            st.session_state["_prod_edit_idx"] = original_idx
                            st.rerun()
                    with c2:
                        if st.button("Delete", key=f"del_{display_idx}"):
                            st.session_state["_prod_delete_idx"] = int(original_idx)
                            st.session_state["_prod_mode"] = "confirm_delete"
                            st.rerun()

    # ---------------- CONFIRM DELETE ----------------
    if st.session_state.get("_prod_mode") == "confirm_delete":
        del_idx = st.session_state.get("_prod_delete_idx")
        if del_idx is not None and 0 <= del_idx < len(df):
            st.warning(f"Confirm delete: {df.iloc[del_idx]['Device']}")
            cdel1, cdel2 = st.columns(2)
            with cdel1:
                if st.button("Yes, Delete"):
                    new_df = df.drop(df.index[del_idx]).reset_index(drop=True)
                    save_products(new_df)
                    st.success("Product deleted")
                    st.session_state.pop("_prod_delete_idx", None)
                    st.session_state.pop("_prod_mode", None)
                    st.rerun()
            with cdel2:
                if st.button("Cancel"):
                    st.session_state.pop("_prod_delete_idx", None)
                    st.session_state.pop("_prod_mode", None)
                    st.rerun()

    # ---------------- PRODUCT CARDS (WORD) ----------------
    st.markdown("---")
    if st.button("Generate Product Cards (Word)"):
        products_df = load_products()
        doc_buf = build_word_cards_document(products_df)
        st.download_button(
            "Download Product Cards (Word)",
            data=doc_buf.getvalue(),
            file_name="product_cards.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # ---------------- IMPORT / EXPORT ----------------
    st.markdown("---")
    st.markdown("<div class='section-title'>Import / Export</div>", unsafe_allow_html=True)

    buf = BytesIO()
    fdf[["Device", "Description", "UnitPrice", "Warranty", "ImageBase64", "ImagePath"]].to_excel(
        buf, index=False
    )
    buf.seek(0)
    st.download_button(
        "Download Products (Excel)",
        data=buf,
        file_name=f"products_export_{datetime.today().strftime('%Y%m%d')}.xlsx",
    )

    up = st.file_uploader("Upload products.xlsx", type=["xlsx"], accept_multiple_files=False)
    if up is not None:
        try:
            imp = pd.read_excel(up)
            for col in ["Device", "Description", "UnitPrice", "Warranty"]:
                if col not in imp.columns:
                    st.error(f"Missing column in uploaded file: {col}")
                    return
            if "ImageBase64" not in imp.columns:
                imp["ImageBase64"] = None
            if "ImagePath" not in imp.columns:
                imp["ImagePath"] = None
            st.warning("This will replace all existing products.")
            ic1, ic2 = st.columns(2)
            with ic1:
                if st.button("Confirm Replace"):
                    imp["Device"] = imp["Device"].apply(proper_case)
                    save_products(
                        imp[["Device", "Description", "UnitPrice", "Warranty", "ImageBase64"]]
                    )
                    st.success("Products replaced from upload.")
                    st.rerun()
            with ic2:
                if st.button("Cancel Import"):
                    st.rerun()
        except Exception as e:
            st.error(f"Failed to read uploaded file: {e}")

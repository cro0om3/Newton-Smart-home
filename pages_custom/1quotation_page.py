import streamlit as st
import pandas as pd
from datetime import datetime
import os
from docx import Document
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
import base64
import tempfile
from streamlit.components.v1 import html as st_html
import convertapi
from pathlib import Path
import sys
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from utils.logger import log_event
from utils.settings import load_settings
 # تم حذف الاستيراد غير المستخدم requests

def proper_case(text):
    if not text:
        return ""
    try:
        return text.title().strip()
    except:
        return text

# Apply the same visual theme used in dashboard_page.py
def _apply_quotation_theme():
    # Now inherits global Invoice theme from main.py
    st.markdown("<style></style>", unsafe_allow_html=True)

def quotation_app():
    _apply_quotation_theme()

        # (Header hero removed to match invoice page)

    # =========================
    # Helper
    # =========================
    def format_phone_input(raw_input):
        digits = ''.join(filter(str.isdigit, raw_input))
        if digits.startswith("0"): digits = digits[1:]
        if digits.startswith("5") and len(digits) == 9:
            return f"+971 {digits[:2]} {digits[2:5]} {digits[5:]}"
        return None

    # =========================
    # UAE Locations
    # =========================
    uae_locations = [
        "Abu Dhabi - Al Shamkha","Abu Dhabi - Al Shawamekh","Abu Dhabi - Khalifa City",
        "Abu Dhabi - Al Bateen","Abu Dhabi - Al Reem Island","Abu Dhabi - Yas Island",
        "Abu Dhabi - Al Mushrif","Abu Dhabi - Al Rawdah","Abu Dhabi - Al Muroor",
        "Abu Dhabi - Baniyas","Abu Dhabi - Mussafah","Abu Dhabi - Al Mafraq",
        "Abu Dhabi - Al Falah","Abu Dhabi - MBZ City","Abu Dhabi - Al Raha",
        "Abu Dhabi - Al Maqtaa","Abu Dhabi - Zayed Port","Abu Dhabi - Saadiyat Island",
        "Al Ain - Al Jimi","Al Ain - Falaj Hazza","Al Ain - Al Maqam",
        "Al Ain - Zakher","Al Ain - Hili","Al Ain - Al Foah","Al Ain - Al Mutaredh",
        "Al Ain - Al Towayya","Al Ain - Al Sarooj","Al Ain - Al Nyadat",
        "Dubai - Marina","Dubai - Downtown","Dubai - Business Bay",
        "Dubai - Jumeirah","Dubai - JBR","Dubai - Al Barsha","Dubai - Mirdif",
        "Dubai - Deira","Dubai - Bur Dubai","Dubai - Silicon Oasis",
        "Dubai - Academic City","Dubai - Arabian Ranches","Dubai - International City",
        "Dubai - Dubai Hills","Dubai - The Springs","Dubai - The Meadows",
        "Dubai - The Greens","Dubai - Palm Jumeirah","Dubai - Al Qusais",
        "Dubai - Al Nahda","Dubai - JVC","Dubai - Damac Hills",
        "Dubai - Discovery Gardens","Dubai - IMPZ","Dubai - Al Warqa",
        "Dubai - Nad Al Sheba",
        "Sharjah - Al Majaz","Sharjah - Al Nahda","Sharjah - Al Taawun",
        "Sharjah - Muwaileh","Sharjah - Al Khan","Sharjah - Al Yarmook",
        "Sharjah - Al Qasimia","Sharjah - Al Fisht","Sharjah - Al Nasserya",
        "Sharjah - Al Goaz","Sharjah - Al Jubail","Sharjah - Maysaloon",
        "Ajman - Al Rashidiya","Ajman - Al Nuaimiya","Ajman - Al Mowaihat",
        "Ajman - Al Rawda","Ajman - Al Jurf","Ajman - Al Hamidiya",
        "Ajman - Al Rumailah","Ajman - Al Bustan","Ajman - City Center",
        "RAK - Al Nakheel","RAK - Al Dhait","RAK - Julph",
        "RAK - Khuzam","RAK - Al Qusaidat","RAK - Seih Al Uraibi",
        "RAK - Al Rams","RAK - Al Mairid","RAK - Mina Al Arab",
        "RAK - Al Hamra Village","RAK - Marjan Island",
        "Fujairah - Al Faseel","Fujairah - Madhab","Fujairah - Dibba",
        "Fujairah - Sakamkam","Fujairah - Mirbah","Fujairah - Al Taween",
        "Fujairah - Kalba","Fujairah - Qidfa","Fujairah - Al Aqah",
        "UAQ - Al Salama","UAQ - Al Haditha","UAQ - Al Raas",
        "UAQ - Al Dar Al Baida","UAQ - Al Khor","UAQ - Al Ramlah",
        "UAQ - Al Maidan","UAQ - Emirates City",
    ]

    # =========================
    # Setup
    # =========================
    try:
        catalog = pd.read_excel("data/products.xlsx")
    except:
        st.error("❌ ERROR: Cannot load product catalog")
        return

    required_cols = ["Device", "Description", "UnitPrice", "Warranty"]
    for col in required_cols:
        if col not in catalog.columns:
            st.error(f"❌ Missing column: {col}")
            return

    # Records helpers (match invoice logic)
    def load_records():
        try:
            df = pd.read_excel("data/records.xlsx")
            df.columns = [c.strip().lower() for c in df.columns]
            return df
        except:
            return pd.DataFrame(columns=[
                "base_id","date","type","number","amount","client_name","phone","location","note"
            ])

    def save_record(rec: dict):
        df = load_records()
        if not df.empty and {"type", "number"}.issubset(df.columns):
            df = df[~((df["type"] == rec.get("type")) & (df["number"] == rec.get("number")))]
        df = pd.concat([df, pd.DataFrame([rec])], ignore_index=True)
        if {"type", "number"}.issubset(df.columns):
            df = df.drop_duplicates(subset=["type", "number"], keep="last")
        df.to_excel("data/records.xlsx", index=False)

    # Customers helpers (auto add from quotation)
    def ensure_customers_file():
        os.makedirs("data", exist_ok=True)
        path = "data/customers.xlsx"
        if not os.path.exists(path):
            cols = [
                "client_name","phone","location","email","status",
                "notes","tags","next_follow_up","assigned_to","last_activity"
            ]
            pd.DataFrame(columns=cols).to_excel(path, index=False)

    def load_customers():
        ensure_customers_file()
        try:
            df = pd.read_excel("data/customers.xlsx")
            df.columns = [c.strip().lower() for c in df.columns]
            return df
        except:
            return pd.DataFrame(columns=[
                "client_name","phone","location","email","status",
                "notes","tags","next_follow_up","assigned_to","last_activity"
            ])

    def save_customers(df: pd.DataFrame):
        os.makedirs("data", exist_ok=True)
        df.to_excel("data/customers.xlsx", index=False)

    def upsert_customer_from_quotation(name: str, phone: str, location: str):
        if not str(name).strip():
            return
        ensure_customers_file()
        cdf = load_customers()
        key = str(name).strip().lower()
        exists = None
        if not cdf.empty and "client_name" in cdf.columns:
            m = cdf["client_name"].astype(str).str.strip().str.lower() == key
            if m.any():
                exists = cdf[m].index[0]
            else:
                # Try phone-based matching when names differ
                cdf_phone = cdf.get("phone", pd.Series([], dtype=str)).apply(lambda x: x if pd.isna(x) else str(x))
                try_phone = phone
                def norm(x):
                    digits = ''.join(filter(str.isdigit, str(x)))
                    if digits.startswith('971') and len(digits) >= 12 and digits[3] == '5':
                        return '0' + digits[3:12]
                    if len(digits) == 9 and digits.startswith('5'):
                        return '0' + digits
                    if len(digits) == 10 and digits.startswith('05'):
                        return digits
                    return digits[-10:]
                if try_phone:
                    m2 = cdf_phone.apply(norm) == norm(try_phone)
                    if m2.any():
                        exists = cdf[m2].index[0]
        if exists is None:
            new_row = {
                "client_name": proper_case(name),
                "phone": phone,
                "location": proper_case(location),
                "email": "",
                "status": "New",
                "notes": "",
                "tags": "",
                "next_follow_up": "",
                "assigned_to": "",
                "last_activity": datetime.today().strftime('%Y-%m-%d'),
            }
            cdf = pd.concat([cdf, pd.DataFrame([new_row])], ignore_index=True)
        else:
            # Update phone/location/last_activity for existing
            cdf.loc[exists, "client_name"] = proper_case(name)
            cdf.loc[exists, "phone"] = phone or cdf.loc[exists, "phone"]
            cdf.loc[exists, "location"] = proper_case(location) or cdf.loc[exists, "location"]
            # Quotation marks engagement start; keep status if set
            cdf.loc[exists, "status"] = cdf.loc[exists, "status"] or "New"
            cdf.loc[exists, "last_activity"] = datetime.today().strftime('%Y-%m-%d')
        save_customers(cdf)

    if "product_table" not in st.session_state:
        st.session_state.product_table = pd.DataFrame(columns=[
            "Item No","Product / Device","Description",
            "Qty","Unit Price (AED)","Line Total (AED)","Warranty (Years)"
        ])

    # =========================
    # CLIENT DETAILS
    # =========================
    st.markdown('<div class="section-title">Quotation Summary</div>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)

    with c1:
        raw_name = st.text_input("Client Name", placeholder="Ahmed Omer", key="quo_client_name")
        client_name = proper_case(raw_name)

        location_selected = st.selectbox("Project Location (UAE)", uae_locations, key="quo_loc")
        client_location = proper_case(location_selected)

        phone_raw = st.text_input("Mobile Number", placeholder="050xxxxxxx", key="quo_phone")
        client_phone = format_phone_input(phone_raw)
        if client_phone:
            st.success(f" {client_phone}")

    with c2:
        today = datetime.today().strftime('%Y%m%d')
        auto_quote = f"QUO-{today}-{len(st.session_state.product_table)+1:03d}"
        quote_no = st.text_input("Quotation No", value=auto_quote, key="quo_no")

        prepared_by = proper_case(st.text_input("Prepared By", value="Mr Bukhari", key="quo_prepared"))
        approved_by = proper_case(st.text_input("Approved By", value="Mr Mohammed", key="quo_approved"))

    # =========================
    # PRODUCTS
    # =========================
    st.markdown("---")
    st.markdown('<div class="section-title">Add Product</div>', unsafe_allow_html=True)

    # Header row
    st.markdown("""
    <div class="product-header">
        <span>Product / Device</span>
        <span>Qty</span>
        <span>Unit Price</span>
        <span>Line Total</span>
        <span>Warranty</span>
        <span>Action</span>
    </div>
    """, unsafe_allow_html=True)

    st.session_state.num_entries = 1

    df = st.session_state.product_table.copy()

    if not df.empty:
        for idx, (i, row) in enumerate(df.iterrows()):
            cols = st.columns([4.5,0.7,1,1,0.7,0.7])

            with cols[0]:
                st.markdown(f"""
                    <div class='added-product-row'>
                        <span style="font-weight:bold;color:rgba(10,132,255,.65);">✓</span>
                        <span style="font-weight:600;color:#1f2937;">{row['Product / Device']}</span>
                    </div>
                """, unsafe_allow_html=True)

            with cols[1]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{int(row['Qty'])}</span></div>", unsafe_allow_html=True)

            with cols[2]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{row['Unit Price (AED)']:.2f}</span></div>", unsafe_allow_html=True)

            with cols[3]:
                st.markdown(
                    f"<div class='added-product-row'><span class='product-value'>AED {row['Line Total (AED)']:.2f}</span></div>",
                    unsafe_allow_html=True
                )

            with cols[4]:
                st.markdown(f"<div class='added-product-row'><span class='product-value'>{int(row['Warranty (Years)'])} yr</span></div>", unsafe_allow_html=True)

            with cols[5]:
                if st.button("❌", key=f"del_q_{i}"):
                    st.session_state.product_table = st.session_state.product_table.drop(i).reset_index(drop=True)
                    st.session_state.product_table["Item No"] = range(1, len(st.session_state.product_table)+1)
                    st.rerun()

    for entry_idx in range(st.session_state.num_entries):
        cols = st.columns([4.5,0.7,1,1,0.7,0.7])

        with cols[0]:
            product = st.selectbox(
                "Product",
                catalog["Device"],
                key=f"prod_entry_{entry_idx}",
                label_visibility="collapsed"
            )
            row = catalog[catalog["Device"] == product].iloc[0]
            desc = row["Description"]

        if f"qty_val_{entry_idx}" not in st.session_state:
            st.session_state[f"qty_val_{entry_idx}"] = 1
        if f"price_val_{entry_idx}" not in st.session_state:
            st.session_state[f"price_val_{entry_idx}"] = float(row["UnitPrice"])
        if f"war_val_{entry_idx}" not in st.session_state:
            st.session_state[f"war_val_{entry_idx}"] = int(row["Warranty"])
        # Sync price and warranty when product changes
        last_key = f"last_prod_{entry_idx}"
        if st.session_state.get(last_key) != product:
            st.session_state[f"price_val_{entry_idx}"] = float(row["UnitPrice"])
            st.session_state[f"war_val_{entry_idx}"] = int(row["Warranty"])
            st.session_state[last_key] = product

        with cols[1]:
            st.number_input(
                "Qty",
                min_value=1,
                step=1,
                value=st.session_state[f"qty_val_{entry_idx}"],
                key=f"qty_val_{entry_idx}",
                label_visibility="collapsed"
            )

        with cols[2]:
            st.number_input(
                "Unit Price (AED)",
                min_value=0.0,
                step=10.0,
                value=st.session_state[f"price_val_{entry_idx}"],
                key=f"price_val_{entry_idx}",
                label_visibility="collapsed"
            )

        qty = st.session_state[f"qty_val_{entry_idx}"]
        price = st.session_state[f"price_val_{entry_idx}"]
        line_price = qty * price

        with cols[3]:
            st.markdown(
                f"<div class='added-product-row'><span class='product-value'>AED {line_price:.2f}</span></div>",
                unsafe_allow_html=True
            )

        with cols[4]:
            st.number_input(
                "Warranty (Years)",
                min_value=0,
                step=1,
                value=st.session_state[f"war_val_{entry_idx}"],
                key=f"war_val_{entry_idx}",
                label_visibility="collapsed"
            )

        warranty = st.session_state[f"war_val_{entry_idx}"]

        with cols[5]:
            if st.button("✅", key=f"add_row_{entry_idx}"):
                new_row = {
                    "Item No": len(st.session_state.product_table) + 1,
                    "Product / Device": product,
                    "Description": desc,
                    "Qty": qty,
                    "Unit Price (AED)": price,
                    "Line Total (AED)": line_price,
                    "Warranty (Years)": warranty,
                }
                st.session_state.product_table = pd.concat(
                    [st.session_state.product_table, pd.DataFrame([new_row])],
                    ignore_index=True
                )
                st.rerun()

    st.markdown("---")

    product_total = st.session_state.product_table["Line Total (AED)"].sum() if not st.session_state.product_table.empty else 0

    # =========================
    # SUMMARY (match invoice)
    # =========================
    st.markdown("---")
    col_left, col_right = st.columns([1, 1])

    with col_left:
        st.markdown("<div class='section-title'>Project Costs</div>", unsafe_allow_html=True)

        # Pull persisted values (so left card reflects right inputs)
        installation_cost_val = st.session_state.get("install_cost_quo_value", 0.0)
        discount_value_val = st.session_state.get("disc_value_quo_value", 0.0)
        discount_percent_val = st.session_state.get("disc_percent_quo_value", 0.0)

        percent_value = (product_total + installation_cost_val) * (discount_percent_val / 100)
        total_discount = percent_value + discount_value_val
        grand_total = (product_total + installation_cost_val) - total_discount

        st.markdown(
            """
            <div style='background:#fff;border:1px solid rgba(0,0,0,.08);border-radius:12px;padding:16px;box-shadow:0 2px 6px rgba(0,0,0,.04);'>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Price (AED)</span>
                    <span style='font-weight:700;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Installation & Operation Devices</span>
                    <span style='font-weight:700;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:10px 0;border-bottom:1px solid rgba(0,0,0,.06);'>
                    <span style='font-weight:600;color:#6e6e73;'>Discount</span>
                    <span style='font-weight:700;color:#1d1d1f;'>-{:,.2f} AED</span>
                </div>
                <div style='display:flex;justify-content:space-between;padding:15px 0;background:rgba(0,0,0,.02);margin-top:8px;border-radius:8px;padding-left:12px;padding-right:12px;'>
                    <span style='font-weight:700;font-size:16px;color:#1d1d1f;'>TOTAL AMOUNT</span>
                    <span style='font-weight:700;font-size:18px;color:#1d1d1f;'>{:,.2f} AED</span>
                </div>
            </div>
            """.format(product_total, installation_cost_val, total_discount, grand_total),
            unsafe_allow_html=True,
        )

    with col_right:
        st.markdown("<div class='section-title'>Installation & Discount</div>", unsafe_allow_html=True)

        installation_cost = st.number_input(
            "Installation & Operation Devices (AED)",
            min_value=0.0,
            step=50.0,
            key="install_cost_quo",
        )
        st.session_state["install_cost_quo_value"] = installation_cost

        st.markdown("<div style='margin-top:12px;'></div>", unsafe_allow_html=True)
        cD1, cD2 = st.columns(2)
        with cD1:
            discount_value = st.number_input("Discount Value (AED)", min_value=0.0, key="disc_value_quo")
            st.session_state["disc_value_quo_value"] = discount_value
        with cD2:
            discount_percent = st.number_input("Discount %", min_value=0.0, max_value=100.0, key="disc_percent_quo")
            st.session_state["disc_percent_quo_value"] = discount_percent

    # =========================
    # EXPORT HELPERS (on-click only)
    # =========================
    def generate_word_file(data: dict) -> BytesIO:
        doc = Document("data/quotation_template.docx")

        # قراءة أبعاد الصور من الإعدادات (سم)
        _s = load_settings()
        _wcm = float(_s.get("quote_product_image_width_cm", 3.49))
        _hcm = float(_s.get("quote_product_image_height_cm", 1.5))

        # خريطة أسماء المنتجات إلى صورة Base64 أو مسارها الأصلي (إن وُجدت)
        image_map = {}
        image_path_map = {}
        try:
            if 'ImageBase64' in catalog.columns:
                image_map = dict(zip(catalog['Device'].astype(str), catalog['ImageBase64']))
            if 'ImagePath' in catalog.columns:
                image_path_map = dict(zip(catalog['Device'].astype(str), catalog['ImagePath']))
        except Exception:
            image_map = {}
            image_path_map = {}

        def insert_image_in_cell(cell, b64_str: str, width_cm: float, height_cm: float, img_path: str = None):
            try:
                bio = None
                if img_path and os.path.exists(img_path):
                    with open(img_path, "rb") as f:
                        bio = BytesIO(f.read())
                elif b64_str and not pd.isna(b64_str):
                    img_bytes = base64.b64decode(b64_str)
                    bio = BytesIO(img_bytes)
                if bio is None:
                    return False
                # تفريغ محتوى الخلية ثم إدراج الصورة في فقرة محاذاة للوسط
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(bio, width=Cm(width_cm), height=Cm(height_cm))
                return True
            except Exception:
                return False

        def format_cell(cell, font_name, font_size, align):
            for paragraph in cell.paragraphs:
                paragraph.alignment = align
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    old = cell.text
                    new = old
                    for key, val in data.items():
                        new = new.replace(key, "" if val is None else str(val))
                    if old != new:
                        cell.text = new
                        format_cell(cell, "Times New Roman (Headings CS)", 10, WD_ALIGN_PARAGRAPH.LEFT)

        # Insert products from session state
        products = st.session_state.product_table.to_dict("records") if "product_table" in st.session_state else []

        target_table = None
        for table in doc.tables:
            try:
                if table.cell(0, 0).text.strip().lower() in ["item no", "item no."]:
                    target_table = table
                    break
            except Exception:
                continue
        if not target_table:
            raise Exception("❌ Product table not found")

        start_row = 1
        last_index = None
        for idx, row in enumerate(target_table.rows):
            if row.cells[0].text.strip().lower() == "last":
                last_index = idx
                break
        if last_index is None:
            raise Exception("❌ 'last' row missing in Word template")

        for i, product in enumerate(products):
            row_index = start_row + i
            if row_index >= last_index:
                break
            row = target_table.rows[row_index]
            row.cells[0].text = str(product.get("Item No", i + 1))
            # إدراج الصورة في عمود المنتج إن وُجدت، وإلا نكتب الاسم نصياً
            prod_name = str(product.get("Product / Device", ""))
            b64_img = image_map.get(prod_name)
            img_path = image_path_map.get(prod_name)
            placed = insert_image_in_cell(row.cells[1], b64_img, _wcm, _hcm, img_path)
            if not placed:
                row.cells[1].text = prod_name
            row.cells[2].text = str(product.get("Description", ""))
            row.cells[3].text = str(product.get("Qty", ""))
            row.cells[4].text = f"{float(product.get('Unit Price (AED)', 0)):,.2f}"
            row.cells[5].text = f"{float(product.get('Line Total (AED)', 0)):,.2f}"
            row.cells[6].text = str(product.get("Warranty (Years)", ""))
            for cell in row.cells:
                format_cell(cell, "Arial MT", 9, WD_ALIGN_PARAGRAPH.CENTER)

        delete_start = start_row + len(products)
        for j in range(last_index - 1, delete_start - 1, -1):
            row = target_table.rows[j]
            target_table._tbl.remove(row._tr)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def convert_to_pdf(word_buffer: BytesIO) -> bytes:
        # Use official ConvertAPI SDK; write temp DOCX then convert
        convertapi.api_credentials = 'HESPs6JNV4IDP62WkWpZe3u7ls8KJA38'
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'quotation.docx')
            with open(docx_path, 'wb') as f:
                f.write(word_buffer.getvalue())
            result = convertapi.convert(
                'pdf',
                {
                    'File': docx_path,
                    'FileName': 'quotation'
                },
                from_format='docx'
            )
            saved = result.save_files(tmpdir)
            if not saved:
                raise Exception('ConvertAPI returned no files')
            pdf_path = saved[0]
            with open(pdf_path, 'rb') as pf:
                return pf.read()

    def _auto_download(data_bytes: bytes, filename: str, mime: str):
        b64 = base64.b64encode(data_bytes).decode('utf-8')
        # Visible fallback link (in case browser blocks auto-download)
        st.markdown(f"If the download doesn't start, click here: [Download {filename}](data:{mime};base64,{b64})", unsafe_allow_html=True)
        st_html(f"""
            <script>
            (function(){{
              const b64 = '{b64}';
              const mime = '{mime}';
              const fname = '{filename}'.replace(/[^\w\-\./]/g,'_');
              const byteChars = atob(b64);
              const byteNumbers = new Array(byteChars.length);
              for (let i = 0; i < byteChars.length; i++) byteNumbers[i] = byteChars.charCodeAt(i);
              const byteArray = new Uint8Array(byteNumbers);
              const blob = new Blob([byteArray], {{type: mime}});
              const url = URL.createObjectURL(blob);
              const a = document.createElement('a');
              a.href = url; a.download = fname; document.body.appendChild(a); a.click();
              setTimeout(()=>{{ URL.revokeObjectURL(url); a.remove(); }}, 1000);
            }})();
            </script>
        """, height=0)

    def _save_export_locally(data_bytes: bytes, filename: str) -> str:
        out_dir = Path('data') / 'exports'
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / filename
        with open(out_path, 'wb') as f:
            f.write(data_bytes)
        return str(out_path)

    st.markdown("---")
    st.markdown('<div class="section-title">Export Quotation</div>', unsafe_allow_html=True)

    # Button colors (blue for Word, red for PDF)
    st.markdown(
        """
        <style>
        div.stButton>button[k="word_action"]{
            background:linear-gradient(145deg,#0a84ff 0%,#1b6cff 100%)!important;color:#fff!important;
            border:1px solid rgba(10,132,255,.35)!important;border-radius:12px!important;
            padding:8px 16px!important;font-weight:700!important;
        }
        div.stButton>button[k="pdf_action"]{
            background:linear-gradient(145deg,#ff3b30 0%,#d70015 100%)!important;color:#fff!important;
            border:1px solid rgba(255,59,48,.35)!important;border-radius:12px!important;
            padding:8px 16px!important;font-weight:700!important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Recalculate using the Installation & Discount values (to mirror invoice)
    product_total = st.session_state.product_table["Line Total (AED)"].sum() if not st.session_state.product_table.empty else 0.0
    installation_cost_val = st.session_state.get("install_cost_quo_value", 0.0)
    discount_value_val = st.session_state.get("disc_value_quo_value", 0.0)
    discount_percent_val = st.session_state.get("disc_percent_quo_value", 0.0)
    percent_value = (product_total + installation_cost_val) * (discount_percent_val / 100)
    total_discount = percent_value + discount_value_val
    grand_total = (product_total + installation_cost_val) - total_discount

    data_to_fill = {
        "{{client_name}}": client_name,
        "{{quote_no}}": quote_no,
        "{{client_location}}": client_location,
        "{{prepared_by}}": prepared_by,
        "{{client_phone}}": client_phone or "N/A",
        "{{approved_by}}": approved_by,
        "{{client_email}}": "N/A",
        # Quotation template keys
        "{{total1}}": f"{product_total:,.2f}",
        "{{installation_cost}}": f"{installation_cost_val:,.2f}",
        "{{Price}}": f"{product_total:,.2f}",
        "{{Total}}": f"{grand_total:,.2f}",
        # Extra keys (no-op if not present in template)
        "{{discount_value}}": f"{discount_value_val:,.2f}",
        "{{discount_percent}}": f"{discount_percent_val:,.0f}",
        "{{total_discount}}": f"{total_discount:,.2f}",
        "{{grand_total}}": f"{grand_total:,.2f}",
    }

    # Always show the two action buttons side-by-side
    b1, b2 = st.columns(2)

    # Simple, invoice-style: pre-render a download_button for Word
    with b1:
        try:
            word_ready = generate_word_file(data_to_fill)
            clicked_word = st.download_button(
                label="Download Word",
                data=word_ready.getvalue(),
                file_name=f"Quotation_{client_name}_{quote_no}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_word_{quote_no}"
            )
            if clicked_word:
                # Save record after user downloads (same behavior as invoice)
                today_id = datetime.today().strftime('%Y%m%d')
                existing = load_records()
                if not existing.empty and "base_id" in existing.columns:
                    same_day = existing[existing.get("base_id", "").astype(str).str.contains(today_id, na=False)]
                    seq = len(same_day) + 1
                else:
                    seq = 1
                base_id = f"{today_id}-{str(seq).zfill(3)}"
                save_record({
                    "base_id": base_id,
                    "date": datetime.today().strftime('%Y-%m-%d'),
                    "type": "q",
                    "number": quote_no,
                    "amount": grand_total,
                    "client_name": client_name,
                    "phone": phone_raw,
                    "location": client_location,
                    "note": ""
                })
                upsert_customer_from_quotation(client_name, phone_raw, client_location)
                # Log quotation creation
                user = st.session_state.get("user", {})
                log_event(user.get("name", "Unknown"), "Quotation", "quotation_created", 
                         f"Client: {client_name}, Amount: {grand_total}")
                st.success(f"✅ Saved quotation to records with base {base_id}")
        except Exception as e:
            st.error(f"❌ Unable to prepare Word file: {e}")

    # PDF: keep a click-to-generate then download button for reliability
    with b2:
        try:
            # Pre-generate PDF bytes for single-click download (may take a moment)
            word_for_pdf = generate_word_file(data_to_fill)
            pdf_ready = convert_to_pdf(word_for_pdf)
            clicked_pdf = st.download_button(
                label="Download PDF",
                data=pdf_ready,
                file_name=f"Quotation_{client_name}_{quote_no}.pdf",
                mime="application/pdf",
                key=f"dl_pdf_{quote_no}"
            )
            if clicked_pdf:
                today_id = datetime.today().strftime('%Y%m%d')
                existing = load_records()
                if not existing.empty and "base_id" in existing.columns:
                    same_day = existing[existing.get("base_id", "").astype(str).str.contains(today_id, na=False)]
                    seq = len(same_day) + 1
                else:
                    seq = 1
                base_id = f"{today_id}-{str(seq).zfill(3)}"
                save_record({
                    "base_id": base_id,
                    "date": datetime.today().strftime('%Y-%m-%d'),
                    "type": "q",
                    "number": quote_no,
                    "amount": grand_total,
                    "client_name": client_name,
                    "phone": phone_raw,
                    "location": client_location,
                    "note": "PDF"
                })
                upsert_customer_from_quotation(client_name, phone_raw, client_location)
                st.success(f"✅ Saved PDF quotation with base {base_id}")
        except Exception as e:
            st.error(f"❌ Unable to prepare PDF: {e}")

